import os
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple
import warnings
warnings.filterwarnings('ignore')


# PDF Processing
import PyPDF2
import pdfplumber

# Word documents
from docx import Document

# Excel
import pandas as pd
from openpyxl import load_workbook


#OCR
try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("OCR libraries not found. OCR functionality will be disabled.")


class DocumentParser:

    def __init__(self, enable_ocr: bool = True):
        self.enable_ocr = enable_ocr and OCR_AVAILABLE
        self.supported_formats = ['.pdf', '.docx', '.xlsx', '.txt', '.csv']

    def parse(self, file_path: str) -> Dict:
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()

        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported format: {extension}")
        
        print(f"Parsing file: {file_path}...")

        # route to appropriate parser
        if extension == '.pdf':
            return self._parse_pdf(file_path)
        elif extension == '.docx':
            return self._parse_docx(file_path)
        elif extension in ['.xlsx', '.xls']:
            return self._parse_excel(file_path)
        elif extension == '.txt':
            return self._parse_text(file_path)
        elif extension == '.csv':
            return self._parse_csv(file_path)
        
    def _parse_pdf(self, file_path: Path) -> Dict:
        result = {
        'filename': file_path.name,
        'format': 'pdf',
        'text': '',
        'pages': [],
        'tables': [],
        'metadata': {},
        'is_scanned': False
        }

        try:
            with pdfplumber.open(file_path) as pdf:
                result['metadata'] = pdf.metadata or {}

                for page_num, page in enumerate(pdf.pages, 1):
                    page_data = {
                        'page_number': page_num,
                        'text': '',
                        'tables': []
                    }

                    #extract text
                    text = page.extract_text()
                    if text:
                        page_data['text'] = text
                        result['text'] += text + "\n\n"

                    # extract tables
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            page_data['tables'].append(df)
                            result['tables'].append({
                                'page': page_num, 'table': df
                            })

                    result['pages'].append(page_data)
            
            if len(result['text'].strip()) < 100:
                result['is_scanned'] = True

                if self.enable_ocr:
                    print(f" Scanned PDF detected. Running OCR...")
                    ocr_text = self._perform_ocr_on_pdf(file_path)
                    result['text'] += ocr_text
                    result['pages'][0]['text'] = ocr_text  # simplistic assignment
            

            print(f"  ✓ Extracted {len(result['pages'])} pages, {len(result['tables'])} tables")
        
        except Exception as e:
            print(f"  ✗ Error parsing PDF: {e}")

            # Fallback: PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    result['metadata'] = pdf_reader.metadata

                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        text = page.extract_text()
                        result['text'] += text + "\n\n"
                        result['pages'].append(
                        {
                            'page_number': page_num,
                            'text': text,
                            'tables': []
                        })
                print(f"  ✓ Extracted {len(result['pages'])} pages using PyPDF2")
            except Exception as e2:
                print(f"  ✗ Error parsing PDF: {e2}")
                raise
        
        return result
    

    def _parse_docx(self, file_path: Path) -> Dict:
        result = {
            'filename': file_path.name,
            'format': 'docx',
            "text": '',
            'paragraphs': [],
            'tables': [],
            'metadata': {}
        }

        try:
            doc = Document(file_path)

            core_props = doc.core_properties
            result['metadata'] = {
                'author': core_props.author,
                'title': core_props.title,
                'created': str(core_props.created),
                'modified': str(core_props.modified)
            }

            for para in doc.paragraphs:
                if para.text.strip():
                    result['paragraphs'].append(para.text)
                    result['text'] += para.text + "\n"

            for table in doc.tables:
                data = []
                for row in table.rows:
                    data.append([cell.text for cell in row.cells])
                if data:
                    df = pd.DataFrame(data[1:], columns=data[0])
                    result['tables'].append(df)
            
            print(f" Extracted {len(result['paragraphs'])} paragraphs, {len(result['tables'])} tables")
        
        except Exception as e:
            print(f" ✗ Error parsing DOCX: {e}")
            raise
        return result
    
    def _parse_excel(self, file_path: Path) -> Dict:
        result = {
            'filename': file_path.name,
            'format': 'xlsx',
            'text': '',
            'sheets': {},
            'tables': []
        }

        try:
            excel_file = pd.ExcelFile(file_path)

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                result['sheets'][sheet_name] = df
                result['tables'].append({
                    'sheet': sheet_name,
                    'data': df
                })

                # Convert to text representation
                result['text'] += f"\n=== Sheet: {sheet_name} ===\n"
                result['text'] += df.to_string + '\n\n'

            print(f"  ✓ Extracted {len(result['sheets'])} sheets")
            
        except Exception as e:
            print(f"  ✗ Error parsing Excel: {e}")
            raise
        
        return result
    

    def _parse_text(self, file_path: Path) -> Dict:
        """Parse plain text files"""
        result = {
            'filename': file_path.name,
            'format': 'txt',
            'text': '',
            'lines': []
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                result['text'] = content
                result['lines'] = content.split('\n')
            
            print(f"  ✓ Extracted {len(result['lines'])} lines")
            
        except Exception as e:
            print(f"  ✗ Error parsing text file: {e}")
            raise
        
        return result
    
    def _parse_csv(self, file_path: Path) -> Dict:
        """Parse CSV files"""
        result = {
            'filename': file_path.name,
            'format': 'csv',
            'text': '',
            'tables': []
        }
        
        try:
            df = pd.read_csv(file_path)
            result['tables'].append({
                'name': file_path.stem,
                'data': df
            })
            result['text'] = df.to_string()
            
            print(f"  ✓ Extracted CSV with {len(df)} rows, {len(df.columns)} columns")
            
        except Exception as e:
            print(f"  ✗ Error parsing CSV: {e}")
            raise
        
        return result
    
    def _ocr_pdf(self, file_path: Path) -> str:
        """
        Perform OCR on scanned PDF
        Requires tesseract installation
        """
        if not OCR_AVAILABLE:
            return "OCR not available. Install pytesseract and pdf2image."
        
        try:
            # Convert PDF to images
            images = convert_from_path(file_path)
            
            text = ""
            for i, image in enumerate(images, 1):
                print(f"    Processing page {i}/{len(images)}...")
                page_text = pytesseract.image_to_string(image)
                text += page_text + '\n\n'
            
            return text
            
        except Exception as e:
            print(f"  ⚠ OCR failed: {e}")
            return ""
    
    def extract_financial_tables(self, parsed_doc: Dict) -> List[pd.DataFrame]:
        """
        Extract tables that look like financial statements
        
        Args:
            parsed_doc: Output from parse()
            
        Returns:
            List of financial tables
        """
        financial_tables = []
        
        # Keywords that indicate financial tables
        financial_keywords = [
            'revenue', 'income', 'profit', 'loss', 'assets', 'liabilities',
            'equity', 'cash flow', 'balance sheet', 'statement', 'quarter',
            'fiscal', 'earnings', 'expenses', 'cost'
        ]
        
        for table_info in parsed_doc.get('tables', []):
            df = table_info.get('data')
            
            if df is None or df.empty:
                continue
            
            # Convert to string for searching
            table_text = df.to_string().lower()
            
            # Check if table contains financial keywords
            if any(keyword in table_text for keyword in financial_keywords):
                financial_tables.append(df)
        
        print(f"  ℹ Found {len(financial_tables)} financial tables")
        return financial_tables
    
    def get_document_summary(self, parsed_doc: Dict) -> Dict:
        """
        Generate summary statistics about parsed document
        """
        summary = {
            'filename': parsed_doc.get('filename'),
            'format': parsed_doc.get('format'),
            'total_text_length': len(parsed_doc.get('text', '')),
            'word_count': len(parsed_doc.get('text', '').split()),
            'table_count': len(parsed_doc.get('tables', [])),
            'page_count': len(parsed_doc.get('pages', [])),
            'is_scanned': parsed_doc.get('is_scanned', False)
        }
        
        return summary


# Example usage and testing
if __name__ == "__main__":
    parser = DocumentParser(enable_ocr=True)
    
    # Example: Parse a PDF
    # result = parser.parse("data/raw/financial_report.pdf")
    # print(result.keys())
    # print(parser.get_document_summary(result))
    
    print("Document Parser Module Ready!")
    print(f"Supported formats: {parser.supported_formats}")
    print(f"OCR enabled: {parser.enable_ocr}")